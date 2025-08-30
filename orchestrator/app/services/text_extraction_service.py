"""
Text extraction service for resume files
"""
import os
import io
from typing import Optional, Dict, Any
from pathlib import Path

# PDF processing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# RTF processing
try:
    import striprtf
    from striprtf.striprtf import rtf_to_text
    RTF_AVAILABLE = True
except ImportError:
    RTF_AVAILABLE = False


class TextExtractionService:
    """Service for extracting text from various file formats"""
    
    def __init__(self):
        self.supported_formats = {
            'txt': self._extract_txt,
            'pdf': self._extract_pdf,
            'docx': self._extract_docx,
            'rtf': self._extract_rtf
        }
    
    def extract_text(self, file_data: bytes, file_type: str) -> Dict[str, Any]:
        """
        Extract text from file data
        
        Args:
            file_data: File content as bytes
            file_type: File extension (pdf, docx, rtf, txt)
            
        Returns:
            Dict with extracted text and metadata
        """
        try:
            file_type = file_type.lower()
            
            if file_type not in self.supported_formats:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            # Check if required library is available
            if file_type == 'pdf' and not PDF_AVAILABLE:
                raise ValueError("PDF processing not available. Install PyPDF2")
            elif file_type == 'docx' and not DOCX_AVAILABLE:
                raise ValueError("DOCX processing not available. Install python-docx")
            elif file_type == 'rtf' and not RTF_AVAILABLE:
                raise ValueError("RTF processing not available. Install striprtf")
            
            # Extract text
            extraction_result = self.supported_formats[file_type](file_data)
            
            return {
                "success": True,
                "text": extraction_result["text"],
                "metadata": extraction_result.get("metadata", {}),
                "file_type": file_type,
                "text_length": len(extraction_result["text"]),
                "error": None
            }
            
        except Exception as e:
            return {
                "success": False,
                "text": "",
                "metadata": {},
                "file_type": file_type,
                "text_length": 0,
                "error": str(e)
            }
    
    def _extract_txt(self, file_data: bytes) -> Dict[str, Any]:
        """Extract text from TXT file"""
        try:
            text = file_data.decode('utf-8')
            return {
                "text": text,
                "metadata": {
                    "encoding": "utf-8",
                    "lines": len(text.split('\n'))
                }
            }
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ['cp1251', 'latin-1', 'iso-8859-1']:
                try:
                    text = file_data.decode(encoding)
                    return {
                        "text": text,
                        "metadata": {
                            "encoding": encoding,
                            "lines": len(text.split('\n'))
                        }
                    }
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Unable to decode text file with any supported encoding")
    
    def _extract_pdf(self, file_data: bytes) -> Dict[str, Any]:
        """Extract text from PDF file"""
        if not PDF_AVAILABLE:
            raise ValueError("PyPDF2 not available")
        
        try:
            pdf_file = io.BytesIO(file_data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            metadata = {
                "pages": len(pdf_reader.pages),
                "title": "",
                "author": "",
                "subject": ""
            }
            
            # Extract text from all pages
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            # Extract metadata
            if pdf_reader.metadata:
                metadata.update({
                    "title": pdf_reader.metadata.get('/Title', ''),
                    "author": pdf_reader.metadata.get('/Author', ''),
                    "subject": pdf_reader.metadata.get('/Subject', '')
                })
            
            return {
                "text": text.strip(),
                "metadata": metadata
            }
            
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_docx(self, file_data: bytes) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            raise ValueError("python-docx not available")
        
        try:
            doc_file = io.BytesIO(file_data)
            doc = Document(doc_file)
            
            text = ""
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections)
            }
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return {
                "text": text.strip(),
                "metadata": metadata
            }
            
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_rtf(self, file_data: bytes) -> Dict[str, Any]:
        """Extract text from RTF file"""
        if not RTF_AVAILABLE:
            raise ValueError("striprtf not available")
        
        try:
            rtf_text = file_data.decode('utf-8')
            plain_text = rtf_to_text(rtf_text)
            
            return {
                "text": plain_text,
                "metadata": {
                    "original_length": len(rtf_text),
                    "extracted_length": len(plain_text)
                }
            }
            
        except Exception as e:
            raise ValueError(f"Failed to extract text from RTF: {str(e)}")
    
    def get_supported_formats(self) -> list:
        """Get list of supported file formats"""
        return list(self.supported_formats.keys())
    
    def is_format_supported(self, file_type: str) -> bool:
        """Check if file format is supported"""
        return file_type.lower() in self.supported_formats
    
    def get_format_requirements(self) -> Dict[str, list]:
        """Get required packages for each format"""
        return {
            "txt": [],
            "pdf": ["PyPDF2"] if not PDF_AVAILABLE else [],
            "docx": ["python-docx"] if not DOCX_AVAILABLE else [],
            "rtf": ["striprtf"] if not RTF_AVAILABLE else []
        }


# Global instance
text_extractor = TextExtractionService()
